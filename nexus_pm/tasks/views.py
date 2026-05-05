import io
import json

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.http import JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone

from accounts.models import User
from products.models import Product

from .decorators import manager_or_admin_required
from .forms import (
    BugReportForm,
    CalendarEventForm,
    CommentForm,
    ProjectForm,
    ProjectModuleForm,
    ReleaseForm,
    TaskForm,
)
from .models import (
    BugReport,
    CalendarEvent,
    KnowledgeBaseNote,
    ModuleMember,
    Notification,
    Project,
    ProjectModule,
    Release,
    Task,
)


def _save_note_as_project_file(note, user):
    """Save a KnowledgeBaseNote as a .md file inside resources/notes/ of its project."""
    if not note.project:
        return
    from django.core.files.base import ContentFile

    from files.models import FileCategory, ProjectFile

    # Ensure Resources/Notes directory tree exists
    resources_cat, _ = FileCategory.objects.get_or_create(
        name="Resources",
        project=note.project,
        parent=None,
        defaults={"created_by": user},
    )
    notes_cat, _ = FileCategory.objects.get_or_create(
        name="Notes",
        project=note.project,
        parent=resources_cat,
        defaults={"created_by": user},
    )

    # Build file content
    md_content = f"# {note.title}\n\n{note.content or ''}"
    filename = f"{note.title[:60].replace(' ', '_')}.md"

    pf = ProjectFile(
        original_name=filename,
        project=note.project,
        category=notes_cat,
        uploaded_by=user,
        description=f"Auto-saved from Knowledge Base note: {note.title}",
        is_public=False,
    )
    pf.file.save(filename, ContentFile(md_content.encode("utf-8")), save=True)


def create_notification(
    recipient, sender, notif_type, title, message, task=None, project=None
):
    if recipient and sender and recipient != sender:
        Notification.objects.create(
            recipient=recipient,
            sender=sender,
            notification_type=notif_type,
            title=title,
            message=message,
            task=task,
            project=project,
        )


def get_visible_tasks_qs(user, tasks_qs):
    if user.is_admin:
        return tasks_qs
    return tasks_qs.filter(
        Q(module__isnull=True)
        | Q(module__members__user=user)
        | Q(project__managers=user)
        | Q(assignees=user)
        | Q(created_by=user)
    ).distinct()


def get_visible_notes_qs(user):
    if user.is_admin:
        return KnowledgeBaseNote.objects.filter(
            Q(project__isnull=False) | Q(project__isnull=True, author=user)
        ).distinct()
    return KnowledgeBaseNote.objects.filter(
        Q(project__managers=user)
        | Q(project__members=user)
        | Q(access_rights__user=user, access_rights__can_view=True)
        | Q(
            project__isnull=True, author=user
        )  # Include general notes authored by the user
    ).distinct()


# ─── DASHBOARD ───────────────────────────────────────────────────────────────


@login_required
def dashboard(request):
    user = request.user
    # Get common data
    today = timezone.now().date()

    if user.is_admin:
        import os
        from django.conf import settings

        db_size = 0
        db_path = settings.DATABASES["default"].get("NAME")
        if db_path and os.path.exists(db_path):
            db_size = os.path.getsize(db_path) / (1024 * 1024)

        projects = Project.objects.all()
        # Active projects for admin: everything not completed or cancelled
        active_projects_count = projects.exclude(status__in=["completed", "cancelled"]).count()
        
        stats = {
            "total_projects": projects.count(),
            "active_projects": active_projects_count,
            "total_users": User.objects.count(),
            "db_size_mb": f"{db_size:.2f}",
            "deletion_reqs": projects.filter(
                Q(deletion_requested_by_admin=True) | Q(deletion_requested_by_pm=True)
            )
            .exclude(deletion_requested_by_admin=True, deletion_requested_by_pm=True)
            .count(),
        }
        return render(
            request,
            "tasks/admin_dashboard.html",
            {"stats": stats, "projects": projects.order_by("-updated_at")[:6]},
        )

    # For PM and regular users
    if user.is_project_manager:
        projects = Project.objects.filter(Q(managers=user) | Q(members=user)).distinct()
    else:
        projects = Project.objects.filter(members=user).distinct()
    
    # Get all visible tasks for the user as base
    all_visible_tasks = get_visible_tasks_qs(user, Task.objects.all())
    
    # My Open Tasks: Assigned to me, visible, and not done
    my_open_tasks_qs = all_visible_tasks.filter(assignees=user).exclude(status="done")
    
    # Overdue Tasks: My open tasks that are past due date
    today = timezone.now().date()
    overdue_tasks_list = [t for t in my_open_tasks_qs if t.due_date and t.due_date < today]
    due_today_list = [t for t in my_open_tasks_qs if t.due_date == today]

    # My Open Bugs: Assigned to me and not resolved/closed/wont_fix
    my_open_bugs_count = BugReport.objects.filter(assignees=user).exclude(
        status__in=["resolved", "closed", "wont_fix"]
    ).count()

    # My bugs list for display (Assigned to me OR Reported by me)
    my_bugs_display = BugReport.objects.filter(Q(assignees=user) | Q(reported_by=user)).exclude(
        status__in=["resolved", "closed", "wont_fix"]
    ).distinct()[:5]

    notifications = Notification.objects.filter(recipient=user, is_read=False)[:5]

    stats = {
        "total_projects": projects.count(),
        "active_projects": projects.exclude(status__in=["completed", "cancelled"]).count(),
        "total_tasks": all_visible_tasks.count(),
        "my_open_tasks": my_open_tasks_qs.count(),
        "overdue_count": len(overdue_tasks_list),
        "completed_tasks": all_visible_tasks.filter(status="done").count(),
        "my_open_bugs": my_open_bugs_count,
    }

    context = {
        "stats": stats,
        "recent_tasks": all_visible_tasks.order_by("-updated_at")[:8],
        "overdue_tasks": overdue_tasks_list[:5],
        "due_today": due_today_list,
        "notifications": notifications,
        "projects": projects.order_by("-updated_at")[:6],
        "my_bugs": my_bugs_display,
    }
    return render(request, "tasks/dashboard.html", context)


# ─── PROJECTS ─────────────────────────────────────────────────────────────────


@login_required
def project_list(request):
    user = request.user
    module_filter = request.GET.get("module", "")
    status_filter = request.GET.get("status", "")
    search = request.GET.get("q", "")
    deletion_requested = request.GET.get("deletion_requested", "")

    if user.is_admin:
        projects = Project.objects.all()
    else:
        projects = Project.objects.filter(Q(managers=user) | Q(members=user)).distinct()

    if deletion_requested:
        projects = projects.filter(Q(deletion_requested_by_admin=True) | Q(deletion_requested_by_pm=True))

    if module_filter:
        projects = projects.filter(module=module_filter)
    if status_filter == "in_progress":
        projects = projects.exclude(status__in=["completed", "cancelled"])
    elif status_filter:
        projects = projects.filter(status=status_filter)
    if search:
        projects = projects.filter(
            Q(name__icontains=search) | Q(description__icontains=search)
        )

    from django.core.paginator import Paginator

    paginator = Paginator(projects.order_by("-created_at"), 10)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)

    return render(
        request,
        "tasks/project_list.html",
        {
            "projects": page_obj,
            "page_obj": page_obj,
            "module_choices": Project.MODULE_CHOICES,
            "status_choices": Project.STATUS_CHOICES,
            "module_filter": module_filter,
            "status_filter": status_filter,
            "search": search,
        },
    )


@login_required
@manager_or_admin_required
def project_create(request):
    form = ProjectForm(request.POST or None, request.FILES or None, user=request.user)
    if request.method == "POST" and form.is_valid():
        project = form.save(commit=False)
        project.created_by = request.user
        project.save()
        form.save_m2m()

        budget_amt = form.cleaned_data.get("budget")
        if budget_amt is not None:
            from finance.models import Budget

            Budget.objects.create(project=project, total_amount=budget_amt)

        # Auto-generate default project folder structure
        from django.core.files.base import ContentFile

        from files.models import FileCategory, ProjectFile

        res_cat = FileCategory.objects.create(
            name="resources", project=project, created_by=request.user
        )
        req_cat = FileCategory.objects.create(
            name="requirements", project=project, created_by=request.user
        )
        rel_cat = FileCategory.objects.create(
            name="releases", project=project, created_by=request.user
        )

        FileCategory.objects.create(
            name="notes", parent=res_cat, project=project, created_by=request.user
        )
        FileCategory.objects.create(
            name="documents", parent=res_cat, project=project, created_by=request.user
        )
        FileCategory.objects.create(
            name="assets", parent=res_cat, project=project, created_by=request.user
        )

        FileCategory.objects.create(
            name="specifications",
            parent=req_cat,
            project=project,
            created_by=request.user,
        )
        test_mgt = FileCategory.objects.create(
            name="test_management",
            parent=req_cat,
            project=project,
            created_by=request.user,
        )
        FileCategory.objects.create(
            name="test_cases", parent=test_mgt, project=project, created_by=request.user
        )
        FileCategory.objects.create(
            name="assessments",
            parent=test_mgt,
            project=project,
            created_by=request.user,
        )

        FileCategory.objects.create(
            name="v1.0", parent=rel_cat, project=project, created_by=request.user
        )
        FileCategory.objects.create(
            name="v2.0", parent=rel_cat, project=project, created_by=request.user
        )

        readme_content = (
            f"# {project.name}\n\n{project.description}\n\nWelcome to your new project."
        )
        readme_file = ProjectFile(
            original_name="README.md",
            project=project,
            uploaded_by=request.user,
            description="Project README",
            is_public=True,
        )
        readme_file.file.save("README.md", ContentFile(readme_content.encode("utf-8")))
        readme_file.save()

        # Notify project manager if assigned
        for pm in project.managers.exclude(pk=request.user.pk):
            create_notification(
                pm,
                request.user,
                "project_update",
                f"You were assigned as Project Manager: {project.name}",
                f'{request.user.display_name} assigned you as the manager of project "{project.name}".',
                project=project,
            )

        # Notify all assigned members
        for member in project.members.all():
            create_notification(
                member,
                request.user,
                "project_update",
                f"You were added to project: {project.name}",
                f'{request.user.display_name} added you as a member of "{project.name}".',
                project=project,
            )
        messages.success(request, f'Project "{project.name}" created successfully.')
        return redirect("tasks:project_detail", pk=project.pk)
    return render(
        request,
        "tasks/project_form.html",
        {"form": form, "title": "New Project", "action": "Create Project"},
    )


@login_required
def project_detail(request, pk):
    project = get_object_or_404(Project, pk=pk)

    if request.user.is_admin:
        messages.error(request, "Admins cannot view inside projects.")
        return redirect("tasks:project_list")

    if not (
        project.members.filter(pk=request.user.pk).exists()
        or project.managers.filter(pk=request.user.pk).exists()
    ):
        messages.error(request, "You do not have access to this project.")
        return redirect("tasks:project_list")

    tasks = (
        get_visible_tasks_qs(request.user, project.tasks.all())
        .select_related("created_by")
        .prefetch_related("assignees")
    )

    # Kanban data
    kanban = {
        "todo": tasks.filter(status="todo"),
        "in_progress": tasks.filter(status="in_progress"),
        "review": tasks.filter(status="review"),
        "done": tasks.filter(status="done"),
        "blocked": tasks.filter(status="blocked"),
    }

    # Filters
    status_filter = request.GET.get("status", "")
    priority_filter = request.GET.get("priority", "")
    assignee_filter = request.GET.get("assignee", "")
    type_filter = request.GET.get("type", "")
    view_mode = request.GET.get("view", "list")

    filtered_tasks = tasks
    if status_filter:
        filtered_tasks = filtered_tasks.filter(status=status_filter)
    if priority_filter:
        filtered_tasks = filtered_tasks.filter(priority=priority_filter)
    if assignee_filter:
        filtered_tasks = filtered_tasks.filter(assignees__id=assignee_filter)
    if type_filter:
        filtered_tasks = filtered_tasks.filter(task_type=type_filter)

    resource_view = request.GET.get("resource_view", "tree")
    repo_cat_id = request.GET.get("repo_cat_id")
    current_repo_cat = None
    if repo_cat_id:
        from files.models import FileCategory
        current_repo_cat = get_object_or_404(FileCategory, pk=repo_cat_id, project=project)

    return render(
        request,
        "tasks/project_detail.html",
        {
            "project": project,
            "tasks": filtered_tasks,
            "kanban": kanban,
            "members": project.members.all(),
            "releases": project.releases.all(),
            "bugs": project.bug_reports.all()[:5],
            "status_choices": Task.STATUS_CHOICES,
            "priority_choices": Task.PRIORITY_CHOICES,
            "type_choices": Task.TYPE_CHOICES,
            "status_filter": status_filter,
            "priority_filter": priority_filter,
            "assignee_filter": assignee_filter,
            "type_filter": type_filter,
            "view_mode": view_mode,
            "resource_view": resource_view,
            "root_categories": project.file_categories.filter(parent=None),
            "current_repo_cat": current_repo_cat,
        },
    )


@login_required
@manager_or_admin_required
def project_edit(request, pk):
    project = get_object_or_404(Project, pk=pk)
    if (
        request.user.is_project_manager
        and not project.managers.filter(pk=request.user.pk).exists()
    ):
        messages.error(
            request, "Only the assigned project manager can edit this project."
        )
        return redirect("tasks:project_detail", pk=project.pk)
    old_members = set(project.members.values_list("pk", flat=True))
    form = ProjectForm(
        request.POST or None, request.FILES or None, instance=project, user=request.user
    )
    if request.method == "POST" and form.is_valid():
        project = form.save()

        budget_amt = form.cleaned_data.get("budget")
        if budget_amt is not None:
            from finance.models import Budget

            budget_obj, created = Budget.objects.get_or_create(project=project)
            budget_obj.total_amount = budget_amt
            budget_obj.save()

        # Notify newly added members
        new_members = set(project.members.values_list("pk", flat=True))
        added_pks = new_members - old_members
        for member in User.objects.filter(pk__in=added_pks):
            create_notification(
                member,
                request.user,
                "project_update",
                f"You were added to project: {project.name}",
                f'{request.user.display_name} added you as a member of "{project.name}".',
                project=project,
            )
        messages.success(request, f'Project "{project.name}" updated.')
        return redirect("tasks:project_detail", pk=project.pk)
    return render(
        request,
        "tasks/project_form.html",
        {
            "form": form,
            "title": f"Edit — {project.name}",
            "action": "Save Changes",
            "project": project,
        },
    )


@login_required
@manager_or_admin_required
def project_members(request, pk):
    """Dedicated page to manage who is on a project."""
    project = get_object_or_404(Project, pk=pk)
    all_users = User.objects.filter(is_active=True).order_by("first_name", "username")
    current_member_ids = set(project.members.values_list("pk", flat=True))

    if request.method == "POST":
        action = request.POST.get("action")
        user_id = request.POST.get("user_id")
        if action and user_id:
            target = get_object_or_404(User, pk=user_id)
            if action == "add":
                project.members.add(target)
                create_notification(
                    target,
                    request.user,
                    "project_update",
                    f"Added to project: {project.name}",
                    f'{request.user.display_name} added you to project "{project.name}".',
                    project=project,
                )
                messages.success(
                    request, f"{target.display_name} added to the project."
                )
            elif action == "remove":
                if project.managers.filter(pk=target.pk).exists():
                    messages.error(
                        request, "Cannot remove the project manager from members."
                    )
                else:
                    project.members.remove(target)
                    messages.success(
                        request, f"{target.display_name} removed from the project."
                    )
        return redirect("tasks:project_members", pk=pk)

    return render(
        request,
        "tasks/project_members.html",
        {
            "project": project,
            "all_users": all_users,
            "current_member_ids": current_member_ids,
        },
    )


@login_required
@manager_or_admin_required
def project_delete(request, pk):
    from datetime import timedelta

    from django.utils import timezone

    project = get_object_or_404(Project, pk=pk)

    if request.method == "POST":
        action = request.POST.get("action", "request_deletion")
        name = project.name

        if action == "request_deletion":
            project.deletion_requested_at = timezone.now()
            if request.user.is_admin:
                project.deletion_requested_by_admin = True
                project.save()
                messages.info(
                    request,
                    f'Project "{name}" deletion requested. Waiting for Project Manager approval.',
                )
                for manager in project.managers.all():
                    create_notification(
                        manager,
                        request.user,
                        "project_update",
                        "Project Deletion Requested",
                        f'Admin {request.user.display_name} has requested to delete project "{name}". Please approve.',
                        project=project,
                    )
            elif (
                request.user.is_project_manager
                and project.managers.filter(pk=request.user.pk).exists()
            ):
                project.deletion_requested_by_pm = True
                project.save()
                messages.info(
                    request,
                    f'Project "{name}" deletion requested. Waiting for Admin approval.',
                )
                for admin in User.objects.filter(role="admin"):
                    create_notification(
                        admin,
                        request.user,
                        "project_update",
                        "Project Deletion Requested",
                        f'PM {request.user.display_name} has requested to delete project "{name}". Please approve.',
                        project=project,
                    )

        elif action == "cancel_deletion":
            if request.user.is_admin and project.deletion_requested_by_admin:
                project.deletion_requested_by_admin = False
                project.deletion_requested_at = None
                project.save()
                messages.info(request, f'Deletion request for "{name}" cancelled.')
            elif (
                request.user.is_project_manager
                and project.managers.filter(pk=request.user.pk).exists()
                and project.deletion_requested_by_pm
            ):
                project.deletion_requested_by_pm = False
                project.deletion_requested_at = None
                project.save()
                messages.info(request, f'Deletion request for "{name}" cancelled.')

        elif action == "approve_deletion":
            if request.user.is_admin and project.deletion_requested_by_pm:
                project.delete()
                messages.success(request, f'Project "{name}" fully deleted.')
                return redirect("tasks:project_list")
            elif (
                request.user.is_project_manager
                and project.managers.filter(pk=request.user.pk).exists()
                and project.deletion_requested_by_admin
            ):
                project.delete()
                messages.success(request, f'Project "{name}" fully deleted.')
                return redirect("tasks:project_list")

        elif action == "force_delete":
            if (
                request.user.is_admin
                and project.deletion_requested_by_admin
                and project.deletion_requested_at
            ):
                if timezone.now() > project.deletion_requested_at + timedelta(days=30):
                    project.delete()
                    messages.success(request, f'Project "{name}" was force deleted.')
                    return redirect("tasks:project_list")
                else:
                    messages.error(
                        request,
                        "You can only force delete after 30 days of requesting.",
                    )

        return redirect("tasks:project_list")

    # Check if 30 days have passed for force delete
    from datetime import timedelta

    from django.utils import timezone

    can_force_delete = False
    if (
        request.user.is_admin
        and project.deletion_requested_by_admin
        and project.deletion_requested_at
    ):
        if timezone.now() > project.deletion_requested_at + timedelta(days=30):
            can_force_delete = True

    return render(
        request,
        "tasks/confirm_delete.html",
        {"obj": project, "obj_type": "Project", "can_force_delete": can_force_delete},
    )


# ─── TASKS ────────────────────────────────────────────────────────────────────


@login_required
def task_list(request):
    user = request.user
    status_filter = request.GET.get("status", "")
    priority_filter = request.GET.get("priority", "")
    search = request.GET.get("q", "")
    my_only = request.GET.get("mine", "")
    project_filter = request.GET.get("project", "")
    overdue_filter = request.GET.get("overdue", "")
    sort = request.GET.get("sort", "-updated_at")
    
    # Validation for sort field to avoid errors
    allowed_sort_fields = ["title", "project__name", "task_type", "priority", "status", "due_date", "updated_at", "-updated_at", "-title", "-project__name", "-task_type", "-priority", "-status", "-due_date"]
    if sort not in allowed_sort_fields:
        sort = "-updated_at"

    if user.is_admin:
        messages.error(request, "Admins do not have access to tasks.")
        return redirect("tasks:dashboard")

    tasks = get_visible_tasks_qs(
        user,
        Task.objects.filter(
            Q(project__members=user) | Q(project__managers=user) | Q(assignees=user)
        ),
    )

    if my_only:
        tasks = tasks.filter(assignees=user)
    if status_filter:
        tasks = tasks.filter(status=status_filter)
    if priority_filter:
        tasks = tasks.filter(priority=priority_filter)
    if project_filter:
        tasks = tasks.filter(project_id=project_filter)
    if overdue_filter:
        from django.utils import timezone
        tasks = tasks.filter(due_date__lt=timezone.now().date()).exclude(status="done")
    if search:
        tasks = tasks.filter(
            Q(title__icontains=search) | Q(description__icontains=search)
        )

    # Projects for filter dropdown
    if user.is_admin:
        projects = Project.objects.all()
    else:
        projects = Project.objects.filter(Q(managers=user) | Q(members=user)).distinct()

    from django.core.paginator import Paginator

    task_qs = (
        tasks.select_related("project")
        .prefetch_related("assignees")
        .order_by(sort)
    )

    # If my_only and no explicit sort, keep the project grouping order
    if my_only and sort == "-updated_at":
        task_qs = (
            tasks.select_related("project")
            .prefetch_related("assignees")
            .order_by("project", "-updated_at")
        )

    paginator = Paginator(task_qs, 10)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)

    # Get current project for theme application
    current_project = None
    if project_filter:
        current_project = Project.objects.filter(id=project_filter).first()

    return render(
        request,
        "tasks/task_list.html",
        {
            "tasks": page_obj,
            "page_obj": page_obj,
            "status_choices": Task.STATUS_CHOICES,
            "priority_choices": Task.PRIORITY_CHOICES,
            "projects": projects,
            "status_filter": status_filter,
            "priority_filter": priority_filter,
            "project_filter": project_filter,
            "search": search,
            "my_only": my_only,
            "my_tasks": my_only,  # alias used by template
            "current_sort": sort,
            "overdue_filter": overdue_filter,
            "project": current_project,
        },
    )


@login_required
@manager_or_admin_required
def task_create(request):
    project_id = request.GET.get("project")
    module_id = request.GET.get("module")
    project = get_object_or_404(Project, pk=project_id) if project_id else None

    initial = {}
    if module_id:
        initial["module"] = get_object_or_404(ProjectModule, pk=module_id)

    form = TaskForm(
        request.POST or None, user=request.user, project=project, initial=initial
    )

    if request.method == "POST" and form.is_valid():
        task = form.save(commit=False)
        task.created_by = request.user
        task.save()
        form.save_m2m()
        for assignee in task.assignees.all():
            if assignee != request.user:
                create_notification(
                    assignee,
                    request.user,
                    "task_assigned",
                    f"New task assigned: {task.title}",
                    f'{request.user.display_name} assigned you a task in "{task.project.name}".',
                    task=task,
                    project=task.project,
                )
        messages.success(request, f'Task "{task.title}" created.')
        return redirect("tasks:task_detail", pk=task.pk)
    return render(
        request,
        "tasks/task_form.html",
        {
            "form": form,
            "title": "New Task",
            "action": "Create Task",
            "project": project,
        },
    )


@login_required
def task_detail(request, pk):
    task = get_object_or_404(Task, pk=pk)

    # Check visibility
    if not get_visible_tasks_qs(request.user, Task.objects.filter(pk=pk)).exists():
        messages.error(request, "You do not have permission to view this task.")
        return redirect("tasks:task_list")

    comments = task.comments.select_related("author").all()
    subtasks = task.subtasks.prefetch_related("assignees").all()
    comment_form = CommentForm()
    project_notes = task.project.kb_notes.all() if task.project else []
    latest_task_files = task.files.filter(versions__isnull=True).select_related("uploaded_by", "parent_file").prefetch_related("versions")

    if request.method == "POST":
        comment_form = CommentForm(request.POST, request.FILES)
        if comment_form.is_valid():
            comment = comment_form.save(commit=False)
            comment.task = task
            comment.author = request.user
            comment.save()

            users_to_notify = set(task.assignees.all())
            if task.created_by:
                users_to_notify.add(task.created_by)
            if task.project:
                users_to_notify.update(task.project.managers.all())

            for user_to_notify in users_to_notify:
                if user_to_notify != request.user:
                    create_notification(
                        user_to_notify,
                        request.user,
                        "comment_added",
                        f"New comment on: {task.title}",
                        f"{request.user.display_name} commented on a task you're involved in.",
                        task=task,
                        project=task.project,
                    )
            messages.success(request, "Comment posted.")
            return redirect("tasks:task_detail", pk=pk)

    return render(
        request,
        "tasks/task_detail.html",
        {
            "task": task,
            "comments": comments,
            "subtasks": subtasks,
            "comment_form": comment_form,
            "project_notes": project_notes,
            "latest_task_files": latest_task_files,
        },
    )


@login_required
def task_edit(request, pk):
    task = get_object_or_404(Task, pk=pk)

    is_pm = task.project.managers.filter(pk=request.user.pk).exists()
    is_assignee = task.assignees.filter(pk=request.user.pk).exists()

    if task.created_by != request.user:
        messages.error(
            request,
            "Only the author of the task can edit it.",
        )
        return redirect("tasks:task_detail", pk=pk)

    old_assignees = set(task.assignees.all())
    old_status = task.status
    form = TaskForm(
        request.POST or None, instance=task, user=request.user, project=task.project
    )
    if request.method == "POST" and form.is_valid():
        new_status = form.cleaned_data.get("status")
        if new_status == "done" and not is_pm and not request.user.is_admin:
            messages.error(
                request,
                "Only Project Managers can mark a task as 'Done'. Please submit 'In Review' instead.",
            )
            return redirect("tasks:task_edit", pk=task.pk)

        task = form.save()

        # Check if release is completed
        if task.status == "done" and old_status != "done" and task.release:
            release = task.release
            if not release.tasks.exclude(status="done").exists():
                for pm in release.project.managers.all():
                    create_notification(
                        pm,
                        request.user,
                        "project_update",
                        f"Release ready for approval: {release.name}",
                        f'All tasks for release "{release.name}" are completed. Please review and approve.',
                        project=release.project,
                    )

        new_assignees = set(task.assignees.all())
        added_assignees = new_assignees - old_assignees
        for assignee in added_assignees:
            if assignee != request.user:
                create_notification(
                    assignee,
                    request.user,
                    "task_assigned",
                    f"Task assigned to you: {task.title}",
                    f"{request.user.display_name} assigned you a task.",
                    task=task,
                    project=task.project,
                )

        if old_status != task.status:
            # If status changed by assignee, notify PM
            if is_assignee and not is_pm:
                for manager in task.project.managers.all():
                    create_notification(
                        manager,
                        request.user,
                        "project_update",
                        f"Task status updated: {task.title}",
                        f'{request.user.display_name} updated the status of task "{task.title}" to {task.get_status_display()}.',
                        task=task,
                        project=task.project,
                    )

            # If status changed by PM (or another user), notify assignees
            if not is_assignee or is_pm:
                for assignee in task.assignees.all():
                    if assignee != request.user:
                        create_notification(
                            assignee,
                            request.user,
                            "project_update",
                            f"Task status updated: {task.title}",
                            f'{request.user.display_name} updated the status of your task "{task.title}" to {task.get_status_display()}.',
                            task=task,
                            project=task.project,
                        )

        messages.success(request, f'Task "{task.title}" updated.')
        return redirect("tasks:task_detail", pk=task.pk)
    return render(
        request,
        "tasks/task_form.html",
        {"form": form, "title": "Edit Task", "action": "Save Changes", "task": task},
    )


@login_required
def task_delete(request, pk):
    task = get_object_or_404(Task, pk=pk)
    project = task.project

    if task.created_by != request.user and not request.user.is_admin:
        messages.error(request, "Only the creator of the task can delete it.")
        return redirect("tasks:task_detail", pk=pk)

    if request.method == "POST":
        task.delete()
        messages.success(request, "Task deleted.")
        return redirect("tasks:project_detail", pk=project.pk)
    return render(
        request, "tasks/confirm_delete.html", {"obj": task, "obj_type": "Task"}
    )


@login_required
def task_update_status(request, pk):
    task = get_object_or_404(Task, pk=pk)
    pms = task.project.managers.all()
    is_pm = request.user in pms

    if not is_pm and not request.user.is_admin:
        return JsonResponse(
            {
                "success": False,
                "message": "Only Project Managers can change the task status.",
            },
            status=403,
        )

    if request.method == "POST":
        try:
            data = json.loads(request.body)
            new_status = data.get("status")

            if new_status in dict(Task.STATUS_CHOICES):
                old_status = task.status
                task.status = new_status
                task.save()

                # Check if release is ready for PM review/completion
                if (
                    task.status in ["review", "done"]
                    and old_status not in ["review", "done"]
                    and task.release
                ):
                    release = task.release
                    if not release.tasks.exclude(
                        status__in=["review", "done"]
                    ).exists():
                        for pm in release.project.managers.all():
                            create_notification(
                                pm,
                                request.user,
                                "project_update",
                                f"Release ready for review & completion: {release.name}",
                                f'All tasks for release "{release.name}" are reviewed or done. Please review and complete the release.',
                                project=release.project,
                            )

                # Notification logic for status change
                for assignee in task.assignees.all():
                    if assignee not in pms and assignee != request.user:
                        create_notification(
                            assignee,
                            request.user,
                            "task_updated",
                            f"Task status changed: {task.title}",
                            f"{request.user.display_name} moved your task to {task.get_status_display()}.",
                            task=task,
                            project=task.project,
                        )

                return JsonResponse(
                    {"success": True, "progress": task.project.progress}
                )
        except Exception:
            pass
    return JsonResponse({"success": False}, status=400)


# ─── NOTIFICATIONS ────────────────────────────────────────────────────────────


@login_required
def notifications(request):
    notifs = Notification.objects.filter(recipient=request.user).select_related(
        "sender", "task", "project"
    )
    status_filter = request.GET.get("status", "unread")
    type_filter = request.GET.get("type", "")

    if request.GET.get("mark_all"):
        notifs.update(is_read=True)
        messages.success(request, "All notifications marked as read.")
        return redirect("tasks:notifications")

    unread_count = notifs.filter(is_read=False).count()

    if status_filter == "unread":
        notifs = notifs.filter(is_read=False)
    elif status_filter == "read":
        notifs = notifs.filter(is_read=True)

    if type_filter:
        notifs = notifs.filter(notification_type=type_filter)

    from django.core.paginator import Paginator

    paginator = Paginator(notifs, 20)
    page_obj = paginator.get_page(request.GET.get("page"))

    return render(
        request,
        "tasks/notifications.html",
        {
            "notifications": page_obj,
            "page_obj": page_obj,
            "unread_count": unread_count,
            "status_filter": status_filter,
            "type_filter": type_filter,
            "type_choices": Notification.TYPE_CHOICES,
        },
    )


@login_required
def notification_read(request, pk):
    notif = get_object_or_404(Notification, pk=pk, recipient=request.user)
    notif.is_read = True
    notif.save()
    if notif.task:
        return redirect("tasks:task_detail", pk=notif.task.pk)
    if notif.project:
        return redirect("tasks:project_detail", pk=notif.project.pk)
    return redirect("tasks:notifications")


# ─── BUGS ─────────────────────────────────────────────────────────────────────


@login_required
def bug_list(request):
    severity_filter = request.GET.get("severity", "")
    status_filter = request.GET.get("status", "")
    project_filter = request.GET.get("project", "")
    assigned_only = request.GET.get("assigned_to_me", "")

    if request.user.is_admin:
        bugs = BugReport.objects.all()
    else:
        bugs = BugReport.objects.filter(
            Q(project__members=request.user)
            | Q(project__managers=request.user)
            | Q(reported_by=request.user)
            | Q(assignees=request.user)
        ).distinct()

    if assigned_only:
        bugs = bugs.filter(assignees=request.user)
    if severity_filter:
        bugs = bugs.filter(severity=severity_filter)
    if status_filter:
        bugs = bugs.filter(status=status_filter)
    if project_filter:
        bugs = bugs.filter(project_id=project_filter)

    if request.user.is_admin:
        projects = Project.objects.all()
    else:
        projects = Project.objects.filter(
            Q(managers=request.user) | Q(members=request.user)
        ).distinct()

    # Get current project for theme application
    current_project = None
    if project_filter:
        current_project = Project.objects.filter(id=project_filter).first()

    return render(
        request,
        "tasks/bug_list.html",
        {
            "bugs": bugs.select_related("project", "reported_by")
            .prefetch_related("assignees")
            .order_by("-created_at"),
            "severity_choices": BugReport.SEVERITY_CHOICES,
            "status_choices": BugReport.STATUS_CHOICES,
            "projects": projects,
            "severity_filter": severity_filter,
            "status_filter": status_filter,
            "project_filter": project_filter,
            "assigned_only": assigned_only,
            "project": current_project,
        },
    )


@login_required
def bug_create(request):
    project_id = request.GET.get("project")
    project = None
    if project_id:
        project = get_object_or_404(Project, pk=project_id)

    form = BugReportForm(request.POST or None, user=request.user, project=project)
    if request.method == "POST" and form.is_valid():
        bug = form.save(commit=False)
        bug.reported_by = request.user
        bug.save()
        form.save_m2m()

        # When anybody reports a bug to a member, add it to assigned person's task list
        if bug.assignees.exists():
            new_task = Task.objects.create(
                title=f"[Bug] {bug.title}",
                description=bug.description,
                project=bug.project,
                task_type="bug",
                status="todo",
                priority=bug.severity,
                created_by=request.user,
            )
            new_task.assignees.set(bug.assignees.all())

        # Notify the person the bug is assigned to
        for assignee in bug.assignees.all():
            if assignee != request.user:
                create_notification(
                    assignee,
                    request.user,
                    "task_assigned",
                    f"Bug assigned to you: {bug.title}",
                    f'{request.user.display_name} assigned you a bug report in "{bug.project.name}": {bug.title}.',
                    project=bug.project,
                )
        messages.success(request, f'Bug "{bug.title}" reported.')
        return redirect("tasks:bug_detail", pk=bug.pk)
    return render(
        request,
        "tasks/bug_form.html",
        {"form": form, "title": "Report a Bug", "action": "Submit Report"},
    )


@login_required
def bug_detail(request, pk):
    bug = get_object_or_404(BugReport, pk=pk)
    return render(request, "tasks/bug_detail.html", {"bug": bug})


@login_required
def bug_edit(request, pk):
    bug = get_object_or_404(BugReport, pk=pk)

    if request.user != bug.reported_by:
        messages.error(request, "Only the author of this bug report can edit it.")
        return redirect("tasks:bug_detail", pk=pk)

    old_assignees = set(bug.assignees.all())
    form = BugReportForm(request.POST or None, instance=bug, user=request.user)
    if request.method == "POST" and form.is_valid():
        bug = form.save()

        new_assignees = set(bug.assignees.all())
        added_assignees = new_assignees - old_assignees

        # Add to assigned person's task list if assigned newly or changed
        if added_assignees:
            new_task = Task.objects.create(
                title=f"[Bug] {bug.title}",
                description=bug.description,
                project=bug.project,
                task_type="bug",
                status="todo",
                priority=bug.severity,
                created_by=request.user,
            )
            new_task.assignees.set(added_assignees)

        # Notify newly assigned person
        for assignee in added_assignees:
            if assignee != request.user:
                create_notification(
                    assignee,
                    request.user,
                    "task_assigned",
                    f"Bug assigned to you: {bug.title}",
                    f'{request.user.display_name} assigned you a bug report in "{bug.project.name}": {bug.title}.',
                    project=bug.project,
                )
        messages.success(request, "Bug report updated.")
        return redirect("tasks:bug_detail", pk=pk)
    return render(
        request,
        "tasks/bug_form.html",
        {
            "form": form,
            "title": "Edit Bug Report",
            "action": "Save Changes",
            "bug": bug,
        },
    )


# ─── CALENDAR ─────────────────────────────────────────────────────────────────


@login_required
def calendar_view(request):
    events = CalendarEvent.objects.filter(
        Q(created_by=request.user) | Q(attendees=request.user)
    ).distinct()
    events_data = [
        {
            "id": e.pk,
            "title": e.title,
            "start": e.start_datetime.isoformat(),
            "end": e.end_datetime.isoformat(),
            "color": e.color,
            "meeting_link": e.meeting_link,
            "meeting_password": e.meeting_password,
        }
        for e in events
    ]

    # Tasks with due dates or deadlines for the user
    my_tasks = Task.objects.filter(
        Q(assignees=request.user) | Q(created_by=request.user),
        Q(due_date__isnull=False) | Q(deadline__isnull=False)
    ).distinct()
    
    for t in my_tasks:
        if t.due_date:
            events_data.append(
                {
                    "id": f"task-due-{t.pk}",
                    "title": f"Task Due: {t.title}",
                    "start": t.due_date.isoformat(),
                    "allDay": True,
                    "color": "#ef4444" if t.is_overdue else "#3b82f6",
                    "url": f"/tasks/{t.pk}/",
                }
            )
        if t.deadline:
            events_data.append(
                {
                    "id": f"task-deadline-{t.pk}",
                    "title": f"Task Deadline: {t.title}",
                    "start": t.deadline.isoformat(),
                    "allDay": True,
                    "color": "#9333ea",  # Purple for deadlines
                    "url": f"/tasks/{t.pk}/",
                }
            )

    # Mix events and tasks in upcoming list
    upcoming_tasks = my_tasks.order_by("due_date")[:5]

    return render(
        request,
        "tasks/calendar.html",
        {
            "events_json": json.dumps(events_data),
            "events": events.order_by("start_datetime")[:10],
            "upcoming_tasks": upcoming_tasks,
            "form": CalendarEventForm(),
        },
    )


@login_required
def event_create(request):
    form = CalendarEventForm(request.POST or None)
    if request.method == "POST" and form.is_valid():
        event = form.save(commit=False)
        event.created_by = request.user
        event.save()
        form.save_m2m()
        
        # Notify project members
        if event.project:
            members = set(event.project.members.all()) | set(event.project.managers.all())
            for member in members:
                if member != request.user:
                    msg = f"A new event '{event.title}' has been scheduled for project {event.project.name}."
                    if event.meeting_link:
                        msg += f" Meeting Link: {event.meeting_link}"
                        if event.meeting_password:
                            msg += f" (Password: {event.meeting_password})"
                    
                    Notification.objects.create(
                        recipient=member,
                        sender=request.user,
                        notification_type="project_update",
                        title=f"New Project Event: {event.title}",
                        message=msg,
                        project=event.project
                    )
        
        messages.success(request, f'Event "{event.title}" created.')
        return redirect("tasks:calendar")
    return render(
        request,
        "tasks/event_form.html",
        {"form": form, "title": "New Event", "action": "Create Event"},
    )


# ─── REPORTS ──────────────────────────────────────────────────────────────────


@login_required
def reports(request):
    if request.user.is_admin:
        messages.info(request, "Reports are not available in admin mode.")
        return redirect("tasks:project_list")
    else:
        projects = Project.objects.filter(
            Q(managers=request.user) | Q(members=request.user)
        ).distinct()
        tasks = get_visible_tasks_qs(
            request.user, Task.objects.filter(project__in=projects)
        )

    task_by_status = {s: tasks.filter(status=s).count() for s, _ in Task.STATUS_CHOICES}
    task_by_priority = {
        p: tasks.filter(priority=p).count() for p, _ in Task.PRIORITY_CHOICES
    }
    proj_by_status = {
        s: projects.filter(status=s).count() for s, _ in Project.STATUS_CHOICES
    }
    proj_by_module = {
        m: projects.filter(module=m).count() for m, _ in Project.MODULE_CHOICES
    }
    overdue_tasks = [
        t
        for t in tasks.select_related("project").prefetch_related("assignees")
        if t.is_overdue
    ]

    # Team workload
    team_workload = []
    for username, _ in User.MODULE_CHOICES:
        members = User.objects.filter(team=username, is_active=True)
        open_tasks = (
            Task.objects.filter(assignees__in=members).exclude(status="done").count()
        )
        team_workload.append(
            {
                "team": username.title(),
                "open_tasks": open_tasks,
                "members": members.count(),
            }
        )

    return render(
        request,
        "tasks/reports.html",
        {
            "projects": projects,
            "tasks": tasks,
            "task_by_status": task_by_status,
            "task_by_priority": task_by_priority,
            "proj_by_status": proj_by_status,
            "proj_by_module": proj_by_module,
            "overdue_tasks": overdue_tasks,
            "total_tasks": tasks.count(),
            "completed_tasks": tasks.filter(status="done").count(),
            "active_projects": projects.filter(status="active").count(),
            "team_workload": team_workload,
        },
    )


# ─── AJAX: Tasks for a project ────────────────────────────────────────────────


@login_required
def tasks_for_project(request):
    """Return JSON list of tasks for a given project (used by file upload & bug forms)."""
    project_id = request.GET.get("project_id")
    if not project_id:
        return JsonResponse({"tasks": []})
    tasks = (
        get_visible_tasks_qs(request.user, Task.objects.filter(project_id=project_id))
        .values("id", "title")
        .order_by("title")
    )
    return JsonResponse({"tasks": list(tasks)})


@login_required
def project_modules_api(request):
    """Return JSON list of modules for a given project (used by task form)."""
    project_id = request.GET.get("project_id")
    if not project_id:
        return JsonResponse({"modules": [], "is_manager": False})
    project = get_object_or_404(Project, pk=project_id)
    is_manager = (
        True
        if project.managers.filter(pk=request.user.pk).exists() or request.user.is_admin
        else False
    )
    modules = (
        ProjectModule.objects.filter(project=project)
        .values("id", "name")
        .order_by("name")
    )
    return JsonResponse({"modules": list(modules), "is_manager": is_manager})


@login_required
def project_requirements_api(request):
    """Return JSON list of requirements for a given project (used by task form)."""
    project_id = request.GET.get("project_id")
    if not project_id:
        return JsonResponse({"requirements": []})
    project = get_object_or_404(Project, pk=project_id)
    # The requirement model must also be imported if not already. I'll just use it directly.
    from .models import Requirement

    requirements = (
        Requirement.objects.filter(project=project)
        .values("id", "name", "req_id")
        .order_by("name")
    )
    return JsonResponse({"requirements": list(requirements)})


@login_required
def project_members_api(request):
    """Return JSON list of members for a given project (used by bug form)."""
    project_id = request.GET.get("project_id")
    if not project_id:
        return JsonResponse({"members": []})
    try:
        project = Project.objects.get(pk=project_id)
    except (Project.DoesNotExist, ValueError):
        return JsonResponse({"members": []})

    member_ids = list(project.members.values_list("pk", flat=True))
    member_ids.extend(project.managers.values_list("pk", flat=True))

    members = User.objects.filter(pk__in=member_ids, is_active=True).order_by(
        "first_name", "username"
    )
    data = [{"id": u.pk, "name": u.display_name} for u in members]
    return JsonResponse({"members": data})


# ─── KNOWLEDGE BASE ──────────────────────────────────────────────────────────


def check_kb_access(kb, user, access_type="view"):
    # If author, always allow
    if kb.author == user:
        return True

    # General notes (no project) are ONLY visible to the author (even admins are excluded)
    if not kb.project:
        return False

    # For project notes, admins and managers have full access
    if user.is_admin:
        return True

    if kb.project and kb.project.managers.filter(pk=user.pk).exists():
        return True

    from files.models import DocumentAccessRight

    explicit = DocumentAccessRight.objects.filter(kb_note=kb, user=user).first()
    if explicit:
        if access_type == "view":
            return explicit.can_view
        if access_type == "edit":
            return explicit.can_edit
        if access_type == "delete":
            return explicit.can_delete

    if access_type != "view":
        return False

    if kb.module:
        return ModuleMember.objects.filter(module=kb.module, user=user).exists()
    elif kb.project:
        return kb.project.members.filter(pk=user.pk).exists()
    return False


@login_required
def kb_overview(request):
    notes = get_visible_notes_qs(request.user)
    q = request.GET.get('q', '')
    project_filter = request.GET.get('project', '')
    author_filter = request.GET.get('author', '')
    
    if q:
        notes = notes.filter(Q(title__icontains=q) | Q(content__icontains=q))
    if project_filter:
        notes = notes.filter(project_id=project_filter)
    if author_filter:
        notes = notes.filter(author_id=author_filter)
        
    from .models import Project
    from accounts.models import User
    
    accessible_projects = Project.objects.filter(
        Q(managers=request.user) | Q(members=request.user)
    ).distinct()
    
    # Authors who have written notes
    authors = User.objects.filter(knowledgebasenote__isnull=False).distinct()
    
    # Get current project for theme application
    current_project = None
    if project_filter:
        current_project = Project.objects.filter(id=project_filter).first()

    return render(request, "tasks/kb_overview.html", {
        "notes": notes, 
        "q": q,
        "projects": accessible_projects,
        "authors": authors,
        "project_filter": project_filter,
        "author_filter": author_filter,
        "project": current_project,
    })


@login_required
def kb_create_global(request):
    if request.user.is_admin:
        projects = Project.objects.all().order_by("name")
    else:
        projects = (
            Project.objects.filter(Q(managers=request.user) | Q(members=request.user))
            .distinct()
            .order_by("name")
        )

    from .forms import KnowledgeBaseNoteForm

    form = KnowledgeBaseNoteForm(request.POST or None)
    if request.method == "POST" and form.is_valid():
        note = form.save(commit=False)
        project_id = request.POST.get("project_id")
        if project_id:
            try:
                note.project = Project.objects.get(pk=project_id)
            except Project.DoesNotExist:
                note.project = None
        note.author = request.user
        note.save()
        _save_note_as_project_file(note, request.user)
        messages.success(request, f'Note "{note.title}" created successfully.')
        return redirect("tasks:kb_overview")

    return render(
        request,
        "tasks/kb_create_global.html",
        {
            "projects": projects,
            "form": form,
        },
    )


@login_required
def kb_list(request, pk):
    project = get_object_or_404(Project, pk=pk)
    if not request.user.is_admin:
        if not (
            project.members.filter(pk=request.user.pk).exists()
            or project.managers.filter(pk=request.user.pk).exists()
        ):
            messages.error(request, "You do not have access to this project.")
            return redirect("tasks:project_list")
    notes = project.kb_notes.all()
    q = request.GET.get('q', '')
    if q:
        notes = notes.filter(Q(title__icontains=q) | Q(content__icontains=q))
    return render(request, "tasks/kb_list.html", {"project": project, "notes": notes, "q": q})


@login_required
def kb_create(request, pk):
    project = get_object_or_404(Project, pk=pk)
    from .models import ProjectModule

    module = None
    module_id = request.GET.get("module")
    if module_id:
        try:
            module = ProjectModule.objects.get(pk=module_id, project=project)
        except ProjectModule.DoesNotExist:
            module = None
    from .forms import KnowledgeBaseNoteForm

    form = KnowledgeBaseNoteForm(request.POST or None)
    if request.method == "POST" and form.is_valid():
        note = form.save(commit=False)
        note.project = project
        if module:
            note.module = module
        note.author = request.user
        note.save()
        _save_note_as_project_file(note, request.user)
        messages.success(
            request, f'Note "{note.title}" created in project "{project.name}".'
        )
        return redirect("tasks:kb_list", pk=project.pk)
    return render(
        request,
        "tasks/kb_form.html",
        {
            "form": form,
            "project": project,
            "module": module,
            "title": "Create Note",
            "action": "Save Note",
        },
    )


@login_required
def kb_detail(request, pk):
    from .models import KnowledgeBaseNote

    note = get_object_or_404(KnowledgeBaseNote, pk=pk)
    project = note.project
    if not check_kb_access(note, request.user, "view"):
        messages.error(request, "You do not have access to this note.")
        return redirect("tasks:project_list")
    return render(request, "tasks/kb_detail.html", {"note": note, "project": project})


@login_required
def kb_edit(request, pk):
    from .models import KnowledgeBaseNote

    note = get_object_or_404(KnowledgeBaseNote, pk=pk)
    project = note.project
    if not check_kb_access(note, request.user, "edit"):
        messages.error(request, "You do not have permission to edit this note.")
        return redirect("tasks:kb_detail", pk=pk)

    from .forms import KnowledgeBaseNoteForm

    form = KnowledgeBaseNoteForm(request.POST or None, instance=note)
    if request.method == "POST" and form.is_valid():
        form.save()
        messages.success(request, "Note updated.")
        return redirect("tasks:kb_detail", pk=pk)
    return render(
        request,
        "tasks/kb_form.html",
        {
            "form": form,
            "project": project,
            "module": note.module,
            "title": "Edit Note",
            "action": "Update Note",
        },
    )


@login_required
def kb_access(request, pk):
    from accounts.models import User
    from files.models import DocumentAccessRight

    from .models import KnowledgeBaseNote

    note = get_object_or_404(KnowledgeBaseNote, pk=pk)
    project = note.project

    if not (
        request.user.is_admin
        or note.author == request.user
        or (project and project.managers.filter(pk=request.user.pk).exists())
    ):
        messages.error(
            request,
            "Only managers, admins, and the author can manage access rights for this KB Note.",
        )
        return redirect("tasks:kb_detail", pk=pk)

    access_rights = DocumentAccessRight.objects.filter(kb_note=note)
    all_users = User.objects.filter(is_active=True)

    if request.method == "POST":
        action = request.POST.get("action")
        if action == "add":
            user_id = request.POST.get("user_id")
            can_view = request.POST.get("can_view") == "on"
            can_edit = request.POST.get("can_edit") == "on"
            can_delete = request.POST.get("can_delete") == "on"
            if user_id:
                target_user = get_object_or_404(User, pk=user_id)
                ar, created = DocumentAccessRight.objects.get_or_create(
                    kb_note=note, user=target_user
                )
                ar.can_view = can_view
                ar.can_edit = can_edit
                ar.can_delete = can_delete
                ar.save()
                messages.success(
                    request, f"Access rights updated for {target_user.display_name}."
                )
        elif action == "remove":
            ar_id = request.POST.get("access_id")
            if ar_id:
                DocumentAccessRight.objects.filter(pk=ar_id).delete()
                messages.success(request, "Access right removed.")
        return redirect("tasks:kb_access", pk=pk)

    return render(
        request,
        "tasks/kb_access.html",
        {"note": note, "access_rights": access_rights, "all_users": all_users},
    )


@login_required
def kb_delete(request, pk):
    from .models import KnowledgeBaseNote

    note = get_object_or_404(KnowledgeBaseNote, pk=pk)
    project = note.project

    if not check_kb_access(note, request.user, "delete"):
        messages.error(request, "You do not have permission to delete this note.")
        return redirect("tasks:kb_detail", pk=pk)

    if request.method == "POST":
        title = note.title
        note.delete()
        messages.success(request, f'Note "{title}" has been deleted.')
        if project:
            return redirect("tasks:kb_list", pk=project.pk)
        else:
            return redirect("tasks:kb_overview")

    return render(
        request, "tasks/kb_confirm_delete.html", {"note": note, "project": project}
    )


# ─── CI/CD & RELEASES ────────────────────────────────────────────────────────
@login_required
def project_cicd(request, pk):
    project = get_object_or_404(Project, pk=pk)
    if not request.user.is_admin:
        if not (
            project.members.filter(pk=request.user.pk).exists()
            or project.managers.filter(pk=request.user.pk).exists()
        ):
            messages.error(request, "You do not have access to this project.")
            return redirect("tasks:project_list")

    pipeline_runs = project.pipeline_runs.all()[:10]
    releases = project.releases.all()

    return render(
        request,
        "tasks/project_cicd.html",
        {"project": project, "pipeline_runs": pipeline_runs, "releases": releases},
    )


# ─── MODULES ──────────────────────────────────────────────────────────────────


@login_required
def module_list(request, pk):
    project = get_object_or_404(Project, pk=pk)
    is_pm = project.managers.filter(pk=request.user.pk).exists()

    if (
        not is_pm
        and not ModuleMember.objects.filter(
            module__project=project, user=request.user
        ).exists()
    ):
        messages.error(
            request, "You do not have access to the modules in this project."
        )
        return redirect("tasks:project_list")

    if is_pm:
        modules = project.modules.all()
    else:
        modules = project.modules.filter(members__user=request.user).distinct()

    return render(
        request, "tasks/module_list.html", {"project": project, "modules": modules}
    )


@login_required
@manager_or_admin_required
def module_create(request, pk):
    project = get_object_or_404(Project, pk=pk)
    form = ProjectModuleForm(request.POST or None)
    if request.method == "POST" and form.is_valid():
        module = form.save(commit=False)
        module.project = project
        module.save()
        messages.success(request, f'Module "{module.name}" created.')
        return redirect("tasks:module_list", pk=project.pk)
    return render(
        request,
        "tasks/module_form.html",
        {"form": form, "project": project, "title": "Create Module"},
    )


@login_required
def requirement_create(request, pk):
    from .forms import RequirementForm
    from .models import Project

    project = get_object_or_404(Project, pk=pk)

    # Must be manager or admin
    if not (
        request.user.is_admin or project.managers.filter(pk=request.user.pk).exists()
    ):
        messages.error(
            request, "Only project managers and admins can create requirements."
        )
        return redirect("tasks:project_detail", pk=project.pk)

    form = RequirementForm(request.POST or None)
    if request.method == "POST" and form.is_valid():
        req = form.save(commit=False)
        req.project = project
        req.save()
        messages.success(request, f'Requirement "{req.name}" created.')
        return redirect(
            f"{reverse('tasks:project_detail', args=[project.pk])}?view=requirements"
        )

    return render(
        request,
        "tasks/requirement_form.html",
        {
            "form": form,
            "project": project,
            "action": "Create Requirement",
            "title": "New Requirement",
        },
    )


@login_required
def requirement_edit(request, pk):
    from .forms import RequirementForm
    from .models import Requirement

    req = get_object_or_404(Requirement, pk=pk)
    project = req.project

    if not (
        request.user.is_admin or project.managers.filter(pk=request.user.pk).exists()
    ):
        messages.error(
            request, "Only project managers and admins can edit requirements."
        )
        return redirect("tasks:project_detail", pk=project.pk)

    form = RequirementForm(request.POST or None, instance=req)
    if request.method == "POST" and form.is_valid():
        form.save()
        messages.success(request, f'Requirement "{req.name}" updated.')
        return redirect(
            f"{reverse('tasks:project_detail', args=[project.pk])}?view=requirements"
        )

    return render(
        request,
        "tasks/requirement_form.html",
        {
            "form": form,
            "project": project,
            "title": "Edit Requirement",
            "action": "Update Requirement",
        },
    )


@login_required
def requirement_delete(request, pk):
    from .models import Requirement

    req = get_object_or_404(Requirement, pk=pk)
    project = req.project

    if not (
        request.user.is_admin or project.managers.filter(pk=request.user.pk).exists()
    ):
        messages.error(
            request, "Only project managers and admins can delete requirements."
        )
        return redirect("tasks:project_detail", pk=project.pk)

    if request.method == "POST":
        name = req.name
        req.delete()
        messages.success(request, f'Requirement "{name}" deleted.')
        return redirect(
            f"{reverse('tasks:project_detail', args=[project.pk])}?view=requirements"
        )

    return render(
        request,
        "tasks/requirement_confirm_delete.html",
        {"requirement": req, "project": project},
    )


@login_required
def module_detail(request, pk):
    module = get_object_or_404(ProjectModule, pk=pk)
    project = module.project

    is_pm = project.managers.filter(pk=request.user.pk).exists()
    is_module_member = module.members.filter(user=request.user).exists()

    if not is_pm and not is_module_member:
        messages.error(request, "You do not have access to this module.")
        return redirect("tasks:project_detail", pk=project.pk)

    members = module.members.all()
    tasks = module.tasks.all()
    files = module.files.all()
    kbs = module.kb_notes.all()
    forum_posts = module.forum_posts.all()

    from .forms import ModuleForumPostForm

    forum_form = ModuleForumPostForm()

    if request.method == "POST":
        forum_form = ModuleForumPostForm(request.POST, request.FILES)
        if forum_form.is_valid():
            post = forum_form.save(commit=False)
            post.author = request.user
            post.module = module
            post.save()

            # Notify members
            module_member_users = [
                m.user for m in module.members.all() if m.user != request.user
            ]
            for pm in project.managers.all():
                if pm not in module_member_users:
                    module_member_users.append(pm)

            for member_user in module_member_users:
                create_notification(
                    member_user,
                    request.user,
                    "project_update",
                    f"New post in module: {module.name}",
                    f'{request.user.display_name} posted in the forum of module "{module.name}".',
                    project=project,
                )

            messages.success(request, "Forum post added.")
            return redirect("tasks:module_detail", pk=module.pk)
    return render(
        request,
        "tasks/module_detail.html",
        {
            "module": module,
            "project": project,
            "members": members,
            "tasks": tasks,
            "files": files,
            "kbs": kbs,
            "forum_posts": forum_posts,
            "forum_form": forum_form,
        },
    )


@login_required
@manager_or_admin_required
def module_edit(request, pk):
    module = get_object_or_404(ProjectModule, pk=pk)
    form = ProjectModuleForm(request.POST or None, instance=module)
    if request.method == "POST" and form.is_valid():
        form.save()
        messages.success(request, f'Module "{module.name}" updated.')
        return redirect("tasks:module_detail", pk=module.pk)
    return render(
        request,
        "tasks/module_form.html",
        {
            "form": form,
            "project": module.project,
            "title": "Edit Module",
            "module": module,
        },
    )


@login_required
@manager_or_admin_required
def module_delete(request, pk):
    module = get_object_or_404(ProjectModule, pk=pk)
    project = module.project
    if request.method == "POST":
        module.delete()
        messages.success(request, "Module deleted.")
        return redirect("tasks:module_list", pk=project.pk)
    return render(
        request,
        "tasks/confirm_delete.html",
        {
            "object": module,
            "title": "Delete Module",
            "cancel_url": reverse("tasks:module_detail", args=[module.pk]),
        },
    )


@login_required
@manager_or_admin_required
def module_members(request, pk):
    module = get_object_or_404(ProjectModule, pk=pk)
    project = module.project
    all_users = User.objects.filter(is_active=True).order_by("first_name", "username")
    current_members = module.members.all()

    if request.method == "POST":
        action = request.POST.get("action")
        user_id = request.POST.get("user_id")
        role = request.POST.get("role", "developer")
        if action and user_id:
            target = get_object_or_404(User, pk=user_id)
            if action == "add":
                ModuleMember.objects.get_or_create(
                    module=module, user=target, defaults={"role": role}
                )
                if (
                    not project.members.filter(pk=target.pk).exists()
                    and not project.managers.filter(pk=target.pk).exists()
                ):
                    project.members.add(target)
                create_notification(
                    target,
                    request.user,
                    "project_update",
                    f"Added to module: {module.name}",
                    f'{request.user.display_name} added you to the module "{module.name}" in project "{project.name}".',
                    project=project,
                )
                messages.success(request, f"{target.display_name} added to the module.")
            elif action == "remove":
                ModuleMember.objects.filter(module=module, user=target).delete()
                messages.success(
                    request, f"{target.display_name} removed from the module."
                )
        return redirect("tasks:module_members", pk=pk)

    return render(
        request,
        "tasks/module_members.html",
        {
            "module": module,
            "project": project,
            "all_users": all_users,
            "current_members": current_members,
        },
    )


# ─── RELEASES ─────────────────────────────────────────────────────────────────


@login_required
def release_list(request, pk):
    project = get_object_or_404(Project, pk=pk)
    # Check basic access
    is_member = (
        project.members.filter(pk=request.user.pk).exists()
        or project.managers.filter(pk=request.user.pk).exists()
        or request.user.is_admin
    )
    is_module_member = ModuleMember.objects.filter(
        module__project=project, user=request.user
    ).exists()

    if not is_member and not is_module_member and not request.user.role == "student":
        messages.error(request, "You do not have access to this project.")
        return redirect("tasks:project_list")

    releases = project.releases.all().order_by("-release_date")
    # End users (e.g. students or non-members who just view phase releases) see only Phase releases
    is_end_user = not (is_member or is_module_member) or request.user.role == "student"
    if is_end_user:
        releases = releases.filter(release_type="phase", is_draft=False)
    else:
        # PM/Admins see all, but we might want to flag drafts
        pass

    latest_release = releases.filter(is_draft=False, is_prerelease=False).first()

    return render(
        request,
        "tasks/release_list.html",
        {"project": project, "releases": releases, "latest_release": latest_release, "is_end_user": is_end_user},
    )


@login_required
@manager_or_admin_required
@login_required
@manager_or_admin_required
def release_create(request, pk=0):
    project = None
    if pk and pk != 0:
        project = get_object_or_404(Project, pk=pk)
    
    form = ReleaseForm(request.POST or None, project=project, user=request.user)
    
    if request.method == "POST" and form.is_valid():
        release = form.save(commit=False)
        if not project:
            project = form.cleaned_data.get('project')
        
        release.project = project
        release.author = request.user
        release.save()
        
        # Handle selected files
        selected_files = form.cleaned_data.get('selected_files')
        if selected_files:
            from files.models import FileCategory
            # Ensure "Releases" category exists
            release_root_cat = FileCategory.objects.filter(name__iexact="Releases", project=project, parent=None).first()
            if not release_root_cat:
                release_root_cat = FileCategory.objects.create(name="Releases", project=project, parent=None, created_by=request.user)
            
            rel_cat, _ = FileCategory.objects.get_or_create(name=release.name, project=project, parent=release_root_cat, defaults={'created_by': request.user})
            
            for f in selected_files:
                f.release = release
                f.category = rel_cat
                f.save()

        messages.success(request, f'Release "{release.name}" created.')
        return redirect("tasks:release_detail", pk=release.pk)
    
    return render(
        request,
        "tasks/release_form.html",
        {
            "form": form,
            "project": project,
            "title": "Create Release",
            "root_categories": project.file_categories.filter(parent=None) if project else []
        },
    )


@login_required
def release_detail(request, pk):
    release = get_object_or_404(Release, pk=pk)
    project = release.project

    # Check access
    is_member = (
        project.members.filter(pk=request.user.pk).exists()
        or project.managers.filter(pk=request.user.pk).exists()
        or request.user.is_admin
    )
    is_module_member = ModuleMember.objects.filter(
        module__project=project, user=request.user
    ).exists()
    is_end_user = not (is_member or is_module_member) or request.user.role == "student"

    if is_end_user and release.release_type != "phase":
        messages.error(request, "You only have access to phase releases.")
        return redirect("tasks:release_list", pk=project.pk)

    tasks = release.tasks.all()
    module_versions = release.module_versions.all()

    if request.method == "POST" and (
        request.user.is_admin
        or project.managers.filter(pk=request.user.pk).exists()
        or request.user == release.author
    ):
        new_status = request.POST.get("status")
        if new_status and new_status in dict(Release.STATUS_CHOICES):
            release.status = new_status
            release.save()
            messages.success(
                request, f"Release status updated to {release.get_status_display()}."
            )
            return redirect("tasks:release_detail", pk=pk)

        if "file" in request.FILES:
            import os
            from files.models import ProjectFile, FileCategory

            # Ensure "Releases" category exists (case-insensitive check)
            release_root_cat = FileCategory.objects.filter(
                name__iexact="Releases",
                project=project,
                parent=None
            ).first()
            
            if not release_root_cat:
                release_root_cat = FileCategory.objects.create(
                    name="Releases",
                    project=project,
                    parent=None,
                    created_by=request.user
                )
            
            # Ensure specific release category exists
            rel_cat, _ = FileCategory.objects.get_or_create(
                name=release.name,
                project=project,
                parent=release_root_cat,
                defaults={'created_by': request.user}
            )

            for uploaded_file in request.FILES.getlist("file"):
                new_f = ProjectFile(
                    file=uploaded_file,
                    original_name=uploaded_file.name,
                    file_size=uploaded_file.size,
                    project=project,
                    release=release,
                    category=rel_cat,
                    uploaded_by=request.user,
                    is_public=True,
                )
                ext = os.path.splitext(uploaded_file.name)[1].lower()
                new_f.extension = ext
                new_f.file_type = ProjectFile.detect_file_type(ext)
                new_f.save()
            
            messages.success(request, f"File(s) added to release and archived under 'Releases/{release.name}'.")
            return redirect("tasks:release_detail", pk=pk)

    kanban = {
        "todo": tasks.filter(status="todo"),
        "in_progress": tasks.filter(status="in_progress"),
        "review": tasks.filter(status="review"),
        "done": tasks.filter(status="done"),
        "blocked": tasks.filter(status="blocked"),
    }

    return render(
        request,
        "tasks/release_detail.html",
        {
            "release": release,
            "project": project,
            "kanban": kanban,
            "tasks": tasks,
            "module_versions": module_versions,
            "project_files": project.files.all().order_by('original_name'),
        },
    )


@login_required
def release_assets_download(request, pk):
    """Download selected files from a project (via release page) as a ZIP."""
    import io
    import os
    import zipfile
    from django.http import HttpResponse
    from files.models import ProjectFile

    release = get_object_or_404(Release, pk=pk)
    project = release.project

    # Access check
    is_member = (
        project.members.filter(pk=request.user.pk).exists()
        or project.managers.filter(pk=request.user.pk).exists()
        or request.user.is_admin
    )
    if not is_member:
        messages.error(request, "Access denied.")
        return redirect("tasks:project_list")

    file_ids = request.POST.getlist("file_ids")
    if not file_ids:
        messages.warning(request, "No files selected.")
        return redirect("tasks:release_detail", pk=pk)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        files = ProjectFile.objects.filter(pk__in=file_ids, project=project)
        files_added = 0
        for pf in files:
            if pf.file and os.path.exists(pf.file.path):
                zip_file.write(pf.file.path, arcname=pf.original_name)
                files_added += 1

    if files_added == 0:
        messages.error(request, "Could not find any of the selected files on disk.")
        return redirect("tasks:release_detail", pk=pk)

    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer.read(), content_type="application/zip")
    response["Content-Disposition"] = f'attachment; filename="assets_{release.tag_name or release.pk}.zip"'
    return response


@login_required
@manager_or_admin_required
def release_edit(request, pk):
    release = get_object_or_404(Release, pk=pk)
    project = release.project
    form = ReleaseForm(request.POST or None, instance=release, project=project)
    if request.method == "POST" and form.is_valid():
        old_name = release.name
        form.save()
        
        # Handle selected files
        selected_files = form.cleaned_data.get('selected_files')
        if selected_files is not None:
            # Clear existing associations to follow selection exactly
            release.direct_files.all().update(release=None)
            
            from files.models import FileCategory
            release_root_cat = FileCategory.objects.filter(name__iexact="Releases", project=project, parent=None).first()
            if not release_root_cat:
                release_root_cat = FileCategory.objects.create(name="Releases", project=project, parent=None, created_by=request.user)
            
            rel_cat, _ = FileCategory.objects.get_or_create(name=release.name, project=project, parent=release_root_cat, defaults={'created_by': request.user})
            
            for f in selected_files:
                f.release = release
                f.category = rel_cat
                f.save()

        if old_name != release.name:
            from files.models import FileCategory
            cat = FileCategory.objects.filter(name=old_name, project=project, parent__name__iexact="Releases").first()
            if cat:
                cat.name = release.name
                cat.save()

        messages.success(request, f'Release "{release.name}" updated.')
        return redirect("tasks:release_detail", pk=release.pk)
    return render(
        request,
        "tasks/release_form.html",
        {
            "form": form,
            "project": project,
            "title": "Edit Release",
            "release": release,
            "root_categories": project.file_categories.filter(parent=None),
        },
    )


@login_required
@manager_or_admin_required
def release_delete(request, pk):
    release = get_object_or_404(Release, pk=pk)
    project = release.project
    if request.method == "POST":
        release.delete()
        messages.success(request, "Release deleted.")
        return redirect("tasks:release_list", pk=project.pk)
    return render(
        request,
        "tasks/confirm_delete.html",
        {
            "object": release,
            "title": "Delete Release",
            "cancel_url": reverse("tasks:release_detail", args=[release.pk]),
        },
    )


@login_required
def release_download(request, pk):
    import io
    import zipfile

    from django.http import HttpResponse

    release = get_object_or_404(Release, pk=pk)
    project = release.project

    is_member = (
        project.members.filter(pk=request.user.pk).exists()
        or project.managers.filter(pk=request.user.pk).exists()
        or request.user.is_admin
    )
    is_module_member = ModuleMember.objects.filter(
        module__project=project, user=request.user
    ).exists()
    is_end_user = not (is_member or is_module_member) or request.user.role == "student"

    if is_end_user and release.release_type != "phase":
        messages.error(request, "You only have access to phase releases.")
        return redirect("tasks:release_list", pk=project.pk)

    zip_buffer = io.BytesIO()
    import os

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        # Include all project files
        project_files = project.files.all().select_related("category")
        files_added = 0
        
        for pf in project_files:
            if pf.file and os.path.exists(pf.file.path):
                # Calculate archive path based on category hierarchy
                parts = []
                cat = pf.category
                while cat:
                    parts.append(cat.name)
                    cat = cat.parent
                
                folder_path = os.path.join(*reversed(parts)) if parts else ""
                archive_name = os.path.join(folder_path, pf.original_name)
                
                # Handle versioning in filename if not version 1
                if pf.version > 1:
                    base, ext = os.path.splitext(pf.original_name)
                    archive_name = os.path.join(folder_path, f"{base}_v{pf.version}{ext}")

                zip_file.write(pf.file.path, arcname=archive_name)
                files_added += 1

        # Also include any specifically attached files that might not be in the project files list
        # (though typically they are)
        for mv in release.module_versions.all():
            if mv.file and mv.file.file and os.path.exists(mv.file.file.path):
                # Check if already added
                pass # Already handled by project.files.all()

        if files_added == 0:
            zip_file.writestr("README.txt", f"No files found for project {project.name}.")

    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer, content_type="application/zip")
    response["Content-Disposition"] = (
        f"attachment; filename={project.name}_{release.name}.zip"
    )
    return response


@login_required
def global_search(request):
    query = request.GET.get('q', '').strip()
    if not query:
        return render(request, "tasks/search_results.html", {"query": query, "results": {}})

    from django.db.models import Q

    from files.models import ProjectFile
    from tasks.models import Project, Task

    # Find tasks
    tasks = Task.objects.filter(
        Q(title__icontains=query) | Q(description__icontains=query)
    ).distinct()

    # Find projects
    projects = Project.objects.filter(
        Q(name__icontains=query) | Q(description__icontains=query)
    ).distinct()

    # Find files
    files = ProjectFile.objects.filter(
        Q(original_name__icontains=query) | Q(description__icontains=query)
    ).distinct()

    # Filter by user permissions
    if not request.user.is_admin:
        tasks = tasks.filter(Q(assignees=request.user) | Q(project__managers=request.user)).distinct()
        projects = projects.filter(Q(members=request.user) | Q(managers=request.user)).distinct()
        # For files, filter by projects the user has access to
        files = files.filter(
            Q(project__members=request.user) | Q(project__managers=request.user)
        ).distinct()

    results = {
        "tasks": tasks[:20],
        "projects": projects[:20],
        "files": files[:20],
    }

    return render(request, "tasks/search_results.html", {
        "query": query,
        "results": results,
    })
@login_required
def inventory_list(request):
    """View to show all products for Project Management users (Read-only)"""
    from django.db.models import Sum

    from stock.models import StockEntry

    search_query = request.GET.get("search", "")
    
    # Start with a base queryset
    products_qs = Product.objects.select_related("category").all().order_by("name")

    if search_query:
        products_qs = products_qs.filter(
            Q(name__icontains=search_query)
            | Q(brand__icontains=search_query)
            | Q(sku__icontains=search_query)
            | Q(serial_number__icontains=search_query)
            | Q(category__name__icontains=search_query)
            | Q(branch__icontains=search_query)
            | Q(rack_number__icontains=search_query)
            | Q(shelf_number__icontains=search_query)
            | Q(description__icontains=search_query)
        )

    from django.core.paginator import Paginator

    # Paginate first to be efficient
    paginator = Paginator(products_qs, 50)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)
    
    # Get the list of products for the current page
    products_list = page_obj.object_list

    # Calculate quantity for each product in the current page
    for product in products_list:
        stock_in = (
            StockEntry.objects.filter(product=product, entry_type="in").aggregate(
                total=Sum("quantity")
            )["total"]
            or 0
        )
        stock_out = (
            StockEntry.objects.filter(product=product, entry_type="out").aggregate(
                total=Sum("quantity")
            )["total"]
            or 0
        )
        from inventory.models import InventoryAdjustment
        adjustments = InventoryAdjustment.objects.filter(product=product).aggregate(
            total=Sum("quantity")
        )["total"] or 0
        product.current_quantity = (stock_in + adjustments) - stock_out

    return render(
        request,
        "tasks/inventory_list.html",
        {
            "products": page_obj, # This is the Page object, iterable in template
            "search_query": search_query,
        },
    )


@login_required
def global_release_list(request):
    """Show all releases across projects grouped by project, with search support."""
    from .models import Release, Project
    from django.db.models import Q
    
    search_query = request.GET.get('q', '')
    
    if request.user.is_admin:
        projects = Project.objects.all()
    else:
        projects = Project.objects.filter(
            Q(members=request.user) | Q(managers=request.user)
        ).distinct()
        
    if search_query:
        projects = projects.filter(
            Q(name__icontains=search_query) | 
            Q(releases__name__icontains=search_query) |
            Q(releases__tag_name__icontains=search_query)
        ).distinct()
        
    projects = projects.order_by("name")
    
    project_releases = []
    for project in projects:
        releases = Release.objects.filter(project=project).order_by("-release_date")
        
        if search_query:
            releases = releases.filter(
                Q(name__icontains=search_query) | 
                Q(tag_name__icontains=search_query)
            )

        if not request.user.is_admin and not project.managers.filter(pk=request.user.pk).exists():
            # Non-managers see only published (non-draft) releases
            releases = releases.filter(Q(is_draft=False) | Q(author=request.user))
            
        if releases.exists():
            project_releases.append({
                'project': project,
                'releases': releases
            })
            
    return render(
        request,
        "tasks/global_release_list.html",
        {
            "project_releases": project_releases,
            "search_query": search_query
        },
    )


# ─── REQUIREMENT REPORT ──────────────────────────────────────────────────────


@login_required
def requirement_report(request, pk):
    """Generate a requirements report for a project in .docx or .md format."""
    from .models import Requirement

    project = get_object_or_404(Project, pk=pk)

    # Access check
    if not (
        request.user.is_admin
        or project.managers.filter(pk=request.user.pk).exists()
        or project.members.filter(pk=request.user.pk).exists()
    ):
        messages.error(request, "You do not have access to this project.")
        return redirect("tasks:project_list")

    export_format = request.GET.get("format", "docx")

    # Fetch requirements with their linked tasks
    requirements = (
        Requirement.objects.filter(project=project)
        .prefetch_related("tasks")
        .order_by("req_id")
    )

    if not requirements.exists():
        messages.warning(request, "No requirements found for this project. Cannot generate report.")
        return redirect(
            f"{reverse('tasks:project_detail', args=[project.pk])}?view=requirements"
        )

    generated_date = timezone.now().strftime("%d %B %Y, %I:%M %p")
    created_date = project.created_at.strftime("%d %B %Y") if project.created_at else "N/A"
    generated_by = request.user.display_name if hasattr(request.user, "display_name") else request.user.username

    # ── MARKDOWN ──────────────────────────────────────────────────────────────
    if export_format == "md":
        lines = []
        lines.append(f"# Requirements Report\n")
        lines.append(f"## Project Information\n")
        lines.append(f"| Field | Value |")
        lines.append(f"|---|---|")
        lines.append(f"| **Project Name** | {project.name} |")
        lines.append(f"| **Project ID** | {project.project_id or project.pk} |")
        lines.append(f"| **Description** | {project.description or 'N/A'} |")
        lines.append(f"| **Created Date** | {created_date} |")
        lines.append(f"| **Generated By** | {generated_by} |")
        lines.append(f"| **Report Generated Date** | {generated_date} |")
        lines.append(f"")
        lines.append(f"---\n")

        lines.append(f"## Requirements Overview\n")
        lines.append(f"| Project Name | Requirement ID | Requirement Name | Description |")
        lines.append(f"|---|---|---|---|")
        for req in requirements:
            req_id = req.req_id or str(req.pk)
            name = req.name.replace("|", "\\|")
            desc = (req.description or "N/A").replace("|", "\\|").replace("\n", " ")
            lines.append(f"| {project.name} | {req_id} | {name} | {desc} |")
        lines.append(f"")
        lines.append(f"---\n")

        lines.append(f"## Requirements with Linked Tasks\n")
        lines.append(f"| Project Name | Requirement ID | Requirement Name | Description | Linked Task ID | Linked Task Name | Linked Task Description |")
        lines.append(f"|---|---|---|---|---|---|---|")
        for req in requirements:
            req_id = req.req_id or str(req.pk)
            req_name = req.name.replace("|", "\\|")
            req_desc = (req.description or "N/A").replace("|", "\\|").replace("\n", " ")
            linked_tasks = req.tasks.all()
            if linked_tasks.exists():
                for task in linked_tasks:
                    task_id = task.task_id or str(task.pk)
                    task_name = task.title.replace("|", "\\|")
                    task_desc = (task.description or "N/A").replace("|", "\\|").replace("\n", " ")
                    lines.append(f"| {project.name} | {req_id} | {req_name} | {req_desc} | {task_id} | {task_name} | {task_desc} |")
            else:
                lines.append(f"| {project.name} | {req_id} | {req_name} | {req_desc} | N/A | N/A | N/A |")

        md_content = "\n".join(lines)
        response = __import__("django.http", fromlist=["HttpResponse"]).HttpResponse(
            md_content, content_type="text/markdown; charset=utf-8"
        )
        response["Content-Disposition"] = 'attachment; filename="requirement_doc.md"'
        return response

    # ── DOCX ──────────────────────────────────────────────────────────────────
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor, Inches
    from django.http import HttpResponse

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Helper: style heading ─────────────────────────────────────────────────
    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF) if level == 1 else RGBColor(0x1D, 0x4E, 0xD8)
        return p

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("Requirements Report", 0)
    title.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"Project: {project.name}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    # ── Section A: Project Information ────────────────────────────────────────
    add_heading("Project Information", level=1)

    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("Project Name", project.name),
        ("Project ID", str(project.project_id or project.pk)),
        ("Description", project.description or "N/A"),
        ("Created Date", created_date),
        ("Generated By", generated_by),
        ("Report Generated Date", generated_date),
    ]
    for i, (label, value) in enumerate(info_data):
        label_cell = info_table.rows[i].cells[0]
        value_cell = info_table.rows[i].cells[1]
        label_cell.width = Inches(2.2)
        value_cell.width = Inches(4.6)
        lp = label_cell.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        vp = value_cell.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(10)

    doc.add_paragraph()

    # ── Section B: Requirements Overview ────────────────────────────────────
    add_heading("Requirements Overview", level=1)

    req_cols = ["Project Name", "Requirement ID", "Requirement Name", "Description"]
    req_table = doc.add_table(rows=1 + requirements.count(), cols=4)
    req_table.style = "Table Grid"

    # Header row
    hdr = req_table.rows[0].cells
    for j, col_name in enumerate(req_cols):
        hp = hdr[j].paragraphs[0]
        hr2 = hp.add_run(col_name)
        hr2.bold = True
        hr2.font.size = Pt(9)
        hr2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Shade header cell
        from docx.oxml import OxmlElement
        tc = hdr[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1E40AF")
        tcPr.append(shd)

    for i, req in enumerate(requirements):
        row = req_table.rows[i + 1].cells
        row[0].paragraphs[0].add_run(project.name).font.size = Pt(9)
        row[1].paragraphs[0].add_run(req.req_id or str(req.pk)).font.size = Pt(9)
        row[2].paragraphs[0].add_run(req.name).font.size = Pt(9)
        row[3].paragraphs[0].add_run(req.description or "N/A").font.size = Pt(9)
        # Zebra striping
        if i % 2 == 1:
            from docx.oxml import OxmlElement
            for cell in row:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EFF6FF")
                tcPr.append(shd)

    doc.add_paragraph()

    # ── Section C: Requirements with Linked Tasks ────────────────────────────
    add_heading("Requirements with Linked Tasks", level=1)

    linked_cols = [
        "Project Name", "Requirement ID", "Requirement Name", "Description",
        "Linked Task ID", "Linked Task Name", "Linked Task Description",
    ]

    # Count total rows needed
    total_rows = 1
    for req in requirements:
        linked_tasks = req.tasks.all()
        total_rows += max(linked_tasks.count(), 1)

    linked_table = doc.add_table(rows=total_rows, cols=7)
    linked_table.style = "Table Grid"

    # Header
    hdr = linked_table.rows[0].cells
    for j, col_name in enumerate(linked_cols):
        hp = hdr[j].paragraphs[0]
        hr3 = hp.add_run(col_name)
        hr3.bold = True
        hr3.font.size = Pt(8)
        hr3.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml import OxmlElement
        tc = hdr[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1E40AF")
        tcPr.append(shd)

    row_idx = 1
    zebra = 0
    for req in requirements:
        req_id_str = req.req_id or str(req.pk)
        linked_tasks = req.tasks.all()
        fill = "EFF6FF" if zebra % 2 == 1 else "FFFFFF"

        if linked_tasks.exists():
            for task in linked_tasks:
                row_cells = linked_table.rows[row_idx].cells
                data = [
                    project.name, req_id_str, req.name, req.description or "N/A",
                    task.task_id or str(task.pk), task.title, task.description or "N/A",
                ]
                for j, val in enumerate(data):
                    p = row_cells[j].paragraphs[0]
                    r = p.add_run(val)
                    r.font.size = Pt(8)
                    if fill != "FFFFFF":
                        from docx.oxml import OxmlElement
                        tc = row_cells[j]._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"), fill)
                        tcPr.append(shd)
                row_idx += 1
        else:
            row_cells = linked_table.rows[row_idx].cells
            data = [project.name, req_id_str, req.name, req.description or "N/A", "N/A", "N/A", "N/A"]
            for j, val in enumerate(data):
                p = row_cells[j].paragraphs[0]
                r = p.add_run(val)
                r.font.size = Pt(8)
                if j >= 4:
                    r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
                if fill != "FFFFFF":
                    from docx.oxml import OxmlElement
                    tc = row_cells[j]._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), fill)
                    tcPr.append(shd)
            row_idx += 1
        zebra += 1

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph(f"Generated by IIA Management System on {generated_date}")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # ── Stream response ───────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = project.name.replace(" ", "_")[:40]
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="requirements_{safe_name}.docx"'
    return response


@login_required
def task_report(request, pk):
    """Generate a task report for a project in .docx or .md format."""
    from .models import Task

    project = get_object_or_404(Project, pk=pk)

    # Access check
    if not (
        request.user.is_admin
        or project.managers.filter(pk=request.user.pk).exists()
        or project.members.filter(pk=request.user.pk).exists()
    ):
        messages.error(request, "You do not have access to this project.")
        return redirect("tasks:project_list")

    export_format = request.GET.get("format", "docx")

    # Fetch tasks for the project
    tasks = (
        Task.objects.filter(project=project)
        .select_related("requirement")
        .order_by("task_id")
    )

    if not tasks.exists():
        messages.warning(request, "No tasks found for this project. Cannot generate report.")
        return redirect(
            f"{reverse('tasks:project_detail', args=[project.pk])}?view=list"
        )

    generated_date = timezone.now().strftime("%d %B %Y, %I:%M %p")
    created_date = project.created_at.strftime("%d %B %Y") if project.created_at else "N/A"
    generated_by = request.user.display_name if hasattr(request.user, "display_name") else request.user.username

    # ── MARKDOWN ──────────────────────────────────────────────────────────────
    if export_format == "md":
        lines = []
        lines.append(f"# Project Task Report\n")
        lines.append(f"## Project Information\n")
        lines.append(f"| Field | Value |")
        lines.append(f"|---|---|")
        lines.append(f"| **Project Name** | {project.name} |")
        lines.append(f"| **Project ID** | {project.project_id or project.pk} |")
        lines.append(f"| **Description** | {project.description or 'N/A'} |")
        lines.append(f"| **Created Date** | {created_date} |")
        lines.append(f"| **Generated By** | {generated_by} |")
        lines.append(f"| **Report Generated Date** | {generated_date} |")
        lines.append(f"")
        lines.append(f"---\n")

        lines.append(f"## Task Details Table\n")
        lines.append(f"| Project Name | Task ID | Task Name | Description |")
        lines.append(f"|---|---|---|---|")
        for task in tasks:
            task_id = task.task_id or str(task.pk)
            title = task.title.replace("|", "\\|")
            desc = (task.description or "N/A").replace("|", "\\|").replace("\n", " ")
            lines.append(f"| {project.name} | {task_id} | {title} | {desc} |")
        lines.append(f"")
        lines.append(f"---\n")

        lines.append(f"## Task-Requirement Mapping Table\n")
        lines.append(f"| Project Name | Task ID | Task Name | Task Description | Linked Requirement ID | Linked Requirement Name | Linked Requirement Description |")
        lines.append(f"|---|---|---|---|---|---|---|")
        for task in tasks:
            task_id = task.task_id or str(task.pk)
            task_name = task.title.replace("|", "\\|")
            task_desc = (task.description or "N/A").replace("|", "\\|").replace("\n", " ")
            
            req = task.requirement
            if req:
                req_id = req.req_id or str(req.pk)
                req_name = req.name.replace("|", "\\|")
                req_desc = (req.description or "N/A").replace("|", "\\|").replace("\n", " ")
                lines.append(f"| {project.name} | {task_id} | {task_name} | {task_desc} | {req_id} | {req_name} | {req_desc} |")
            else:
                lines.append(f"| {project.name} | {task_id} | {task_name} | {task_desc} | N/A | N/A | N/A |")

        md_content = "\n".join(lines)
        response = __import__("django.http", fromlist=["HttpResponse"]).HttpResponse(
            md_content, content_type="text/markdown; charset=utf-8"
        )
        response["Content-Disposition"] = 'attachment; filename="task_report.md"'
        return response

    # ── DOCX ──────────────────────────────────────────────────────────────────
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor, Inches
    from django.http import HttpResponse

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF) if level == 1 else RGBColor(0x1D, 0x4E, 0xD8)
        return p

    # Title
    title = doc.add_heading("Project Task Report", 0)
    title.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"Project: {project.name}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    # Section 1: Project Information
    add_heading("Project Information", level=1)
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("Project Name", project.name),
        ("Project ID", str(project.project_id or project.pk)),
        ("Description", project.description or "N/A"),
        ("Created Date", created_date),
        ("Generated By", generated_by),
        ("Report Generated Date", generated_date),
    ]
    for i, (label, value) in enumerate(info_data):
        label_cell = info_table.rows[i].cells[0]
        value_cell = info_table.rows[i].cells[1]
        label_cell.width = Inches(2.0)
        value_cell.width = Inches(4.5)
        lp = label_cell.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        vp = value_cell.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(10)

    doc.add_paragraph()

    # Section 2: Task Details Table
    add_heading("Task Details Table", level=1)
    task_cols = ["Project Name", "Task ID", "Task Name", "Description"]
    task_table = doc.add_table(rows=1 + tasks.count(), cols=4)
    task_table.style = "Table Grid"

    # Header
    hdr = task_table.rows[0].cells
    for j, col_name in enumerate(task_cols):
        hp = hdr[j].paragraphs[0]
        hr2 = hp.add_run(col_name)
        hr2.bold = True
        hr2.font.size = Pt(9)
        hr2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml import OxmlElement
        tc = hdr[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1E40AF")
        tcPr.append(shd)

    for i, task in enumerate(tasks):
        row = task_table.rows[i + 1].cells
        row[0].paragraphs[0].add_run(project.name).font.size = Pt(9)
        row[1].paragraphs[0].add_run(task.task_id or str(task.pk)).font.size = Pt(9)
        row[2].paragraphs[0].add_run(task.title).font.size = Pt(9)
        row[3].paragraphs[0].add_run(task.description or "N/A").font.size = Pt(9)
        if i % 2 == 1:
            from docx.oxml import OxmlElement
            for cell in row:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EFF6FF")
                tcPr.append(shd)

    doc.add_paragraph()

    # Section 3: Task-Requirement Mapping Table
    add_heading("Task-Requirement Mapping Table", level=1)
    mapping_cols = ["Project Name", "Task ID", "Task Name", "Task Description", "Req ID", "Req Name", "Req Description"]
    
    mapping_table = doc.add_table(rows=1 + tasks.count(), cols=7)
    mapping_table.style = "Table Grid"

    # Header
    hdr = mapping_table.rows[0].cells
    for j, col_name in enumerate(mapping_cols):
        hp = hdr[j].paragraphs[0]
        hr3 = hp.add_run(col_name)
        hr3.bold = True
        hr3.font.size = Pt(8)
        hr3.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml import OxmlElement
        tc = hdr[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1E40AF")
        tcPr.append(shd)

    for i, task in enumerate(tasks):
        task_id_str = f"{project.name} - {task.task_id or str(task.pk)}"
        req = task.requirement
        fill = "EFF6FF" if i % 2 == 1 else "FFFFFF"

        row_cells = mapping_table.rows[i + 1].cells
        if req:
            data = [
                project.name, task_id_str, task.title, task.description or "N/A",
                f"{project.name} - {req.req_id or str(req.pk)}", req.name, req.description or "N/A"
            ]
        else:
            data = [project.name, task_id_str, task.title, task.description or "N/A", "N/A", "N/A", "N/A"]

        for j, val in enumerate(data):
            p = row_cells[j].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8)
            if not req and j >= 4:
                r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
            
            if fill != "FFFFFF":
                from docx.oxml import OxmlElement
                tc = row_cells[j]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), fill)
                tcPr.append(shd)

    doc.add_paragraph()
    footer_p = doc.add_paragraph(f"Generated by IIA Management System on {generated_date}")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = project.name.replace(" ", "_")[:40]
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="task_report_{safe_name}.docx"'
    return response
